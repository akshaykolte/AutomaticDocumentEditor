import datetime
import os
import docx2txt
import docx
from docxtpl import *
from watchdog.observers import Observer
from watchdog.events import LoggingEventHandler
import logging
from random import randrange as rr
import ntpath

from utilities import *
from read_excel import *

def read_pc_doc_file(doc_file_name):
    try:
        text = docx2txt.process(doc_file_name)
        print 'Doc file %s found'%(doc_file_name)
    except:
        print 'No Doc file %s found'%(doc_file_name)
        return []

    index = 0
    start_text = 'The said goods may be allowed to be cleared at NIL rate of duty in terms of Notification 52/2003-Cus dated 31.03.2006, under intimation to the undersigned.'
    end_text = 'OFFICE OF THE SUPERINTENDENT OF CENTRAL EXCISE, WAGHOLI'
    address_list = []
    while index < len(text):
        start = text.find(start_text, index) + len(start_text)
        end = text.find(end_text, start)
        if end >=0:
            address_text = text[start:end]
            address_list.append(text[start:end].strip().replace('\n\n', '\n'))
            index = end
        else:
            address_text = text[start:]
            address_list.append(text[start:].strip().replace('\n\n', '\n'))
            index = len(text)


    index = 0
    start_text = 'Balance in B-17 Bond Rs.'
    end_text = 'The said goods may be allowed to be cleared at NIL rate of duty in terms of Notification 52/2003-Cus dated 31.03.2006, under intimation to the undersigned.'
    balance_list = []
    # print text
    while index < len(text):
        start = text.find(start_text, index) + len(start_text)
        if start < len(start_text):
            break
        end = text.find(end_text, start)
        balance_text = text[start:end].strip()
        balance_list.append(balance_text)
        index = end

    index = 0
    start_text = 'This registration authorizes them to obtain/clear material from'
    end_text = 'for re'
    short_address_list = []
    # print text
    while index < len(text):
        start = text.find(start_text, index) + len(start_text)
        if start < len(start_text):
            break
        end = text.find(end_text, start)
        short_address_text = text[start:end].strip()
        short_address_list.append(short_address_text)
        index = end

    document = docx.Document(doc_file_name)
    table_list = []
    tables = document.tables
    for i in range(1,len(tables), 2):
        table = tables[i]
        table_list.append({})
        table_list[-1]['sr_no'] = get_text_from_cell(table.rows[1].cells[0])
        table_list[-1]['description_of_goods'] = get_text_from_cell(table.rows[1].cells[1])
        table_list[-1]['qty'] = get_text_from_cell(table.rows[1].cells[2])
        table_list[-1]['cif_value'] = get_text_from_cell(table.rows[1].cells[3])
        table_list[-1]['total'] = get_text_from_cell(table.rows[2].cells[3])
        table_list[-1]['address'] = address_list[i/2]
        table_list[-1]['balance'] = balance_list[i/2]
        table_list[-1]['short_address'] = short_address_list[i/2]

    return table_list

def read_ic_doc_file(doc_file_name):
    table_dict = {}
    table_dict['capital_goods_limit'] = '44,25,00,000.00'
    table_dict['running_balance'] = ''
    table_dict['email_id_and_phone_number'] = 'Suresh Shinde, Suresh.shinde@axa-abs.co.in,  Assistant Manager, 9890858845'
    try:
        text = docx2txt.process(doc_file_name)
        print 'Doc file %s found'%(doc_file_name)
    except:
        print 'No Doc file %s found'%(doc_file_name)
        return table_dict

    index = 0
    start_text = 'Initial Approved Imported Capital Goods Limit'
    end_text = 'Running Balance (Excluding this import)'
    index = text.find(start_text)
    index = text.find('Rs. ', index) + 4
    end_index = text.find(end_text)
    capital_goods_limit = text[index:end_index].strip()
    table_dict['capital_goods_limit'] = capital_goods_limit

    index = text.find(end_text)
    index = text.find('Rs. ', index) + 4
    end_index = text.find('Undertaking', index)
    running_balance = text[index:end_index].strip()
    table_dict['running_balance'] = running_balance

    document = docx.Document(doc_file_name)
    table = document.tables[0]
    table_dict['email_id_and_phone_number'] = get_text_from_cell(table.rows[16].cells[2])
    return table_dict

def save_to_pc_doc_file(doc_file_name, unit_details, table_list):
    tpl=DocxTemplate(PC_DOCX_TEMPLATE)
    context = {}
    context['ic_date'] = unit_details['ic_date']
    context['pc_date'] = unit_details['pc_date']
    context['sr_no'] = table_list['sr_no']
    context['description_of_goods'] = table_list['description_of_goods']
    context['qty'] = table_list['qty']
    context['cif_value'] = table_list['cif_value']
    context['total'] = table_list['total']
    context['debit'] = unit_details['debit-b17']
    context['address'] = table_list['address']
    context['short_address'] = table_list['short_address']
    context['balance'] = table_list['balance']
    context['certificate_number'] = unit_details['certificate_number']
    context['supplier_name'] = unit_details['name_of_supplier']
    context['country'] = unit_details['country_name']
    context['proforma_invoice_number'] = unit_details['invoice_number_and_date']
    tpl.render(context)
    tpl.save(doc_file_name)

def save_to_rwc_doc_file(doc_file_name, unit_details, table_list):
    tpl=DocxTemplate(RWC_DOCX_TEMPLATE)
    context = {}
    context['address'] = table_list['address']
    context['procurement_certificate'] = unit_details['procurement_certificate_number']
    context['pc_date'] = unit_details['pc_date']
    context['particulars'] = table_list['particulars']
    context['ta_number'] = unit_details['ta_number']
    context['ta_date'] = unit_details['ta_date']
    context['boe_number_and_date'] = unit_details['boe_number_and_date']
    context['ta_date'] = unit_details['ta_date']
    context['supplier_name'] = unit_details['name_of_supplier']
    context['country'] = unit_details['country_name']
    context['description_of_goods'] = table_list['description_of_goods']
    context['gross_weight'] = table_list['gross_weight']
    context['assessable_value'] = unit_details['assessable_value']
    context['total_duty_foregone'] = unit_details['duty_foregone']
    context['goods_receipt_date'] = unit_details['bond_page_date']
    context['date_of_rewarehousing'] = table_list['date_of_rewarehousing']
    context['bond_register_page_number'] = unit_details['bond_page_number']
    context['shortage'] = table_list['shortage']
    tpl.render(context)
    tpl.save(doc_file_name)

def save_to_ic_doc_file(doc_file_name, unit_details):
    tpl=DocxTemplate(IC_DOCX_TEMPLATE)
    context = {}

    context['SerialNo'] = unit_details['SerialNo']
    context['ClientName'] = unit_details['ClientName']
    context['SubBranch'] = unit_details['SubBranch']
    context['ClientAddress'] = unit_details['ClientAddress']
    context['LOPNoAndDate'] = unit_details['LOPNoAndDate']
    context['LegalUndertakingDate'] = unit_details['LegalUndertakingDate']
    context['ExtensionofLOPNoAndDate'] = unit_details['ExtensionofLOPNoAndDate']
    context['CustomBondingLicenseNoAndDate'] = unit_details['CustomBondingLicenseNoAndDate']
    context['LicenseValidityDate'] = unit_details['LicenseValidityDate']
    context['InvoiceNoAndDate'] = unit_details['InvoiceNoAndDate']
    context['BroadDescription'] = unit_details['BroadDescription']
    context['PurposeOfUtilizationOfTheItem'] = unit_details['PurposeOfUtilizationOfTheItem']
    context['NewUsedEquipment'] = unit_details['NewUsedEquipment']
    context['AppliedItemBroaderCategory'] = unit_details['AppliedItemBroaderCategory']
    context['NameOfSupplier'] = unit_details['NameOfSupplier']
    context['CountryName'] = unit_details['CountryName']
    context['Incoterm'] = unit_details['Incoterm']
    context['CIFValue'] = 'CIF/USD: ' + ((unit_details['CIFValue']))
    context['Currency'] = unit_details['Currency']
    context['ClientContactPersonNameEmailAndContactNo'] = unit_details['ClientContactPersonNameEmailAndContactNo'].strip()
    context['CGApprovedAmount'] = unit_details['CGApprovedAmount']
    context['RunningCGBalance'] = unit_details['RunningCGBalance']
    context['ICApplicationDate'] = unit_details['ICApplicationDate']
    context['Place'] = unit_details['Place']

    tpl.render(context)
    tpl.save(doc_file_name)

def update_pc_file(unit_details_dict):

    base_unit_director = BASE_DIR
    
    # create directory for unit, if not exists
    if not os.path.exists(base_unit_director):
        os.makedirs(base_unit_director)

    pc_folder = base_unit_director + 'PC' + SLASH

    # create IC directory if not exists
    if not os.path.exists(pc_folder):
        os.makedirs(pc_folder)
    
    for unit_details in unit_details_dict:

        print unit_details
        file_name=pc_folder
        print  "+++++++++++++"
        print unit_details['ClientName']
        print unit_details['SubBranch'].strip()

        if not os.path.exists(pc_folder+unit_details['ClientName']+SLASH):
            os.makedirs(pc_folder+unit_details['ClientName']+SLASH)
        
        file_name = pc_folder+unit_details['ClientName']+SLASH
        if unit_details['SubBranch'].strip() != '':
            file_name = pc_folder+unit_details['ClientName']+SLASH+unit_details['SubBranch']+SLASH
            if not os.path.exists(pc_folder+unit_details['ClientName']+SLASH+unit_details['SubBranch']+SLASH):
                os.makedirs(pc_folder+unit_details['ClientName']+SLASH+unit_details['SubBranch']+SLASH)               
        else:
            file_name = pc_folder+unit_details['ClientName']+SLASH
      
        file_name = file_name + 'PC Application '

        file_name += (unit_details['ICApplicationDate'] + '({})'.format(unit_details['SerialNo'])).strip() + '.docx'
        print file_name            
        print "END+============================="       
        #save_to_pc_doc_file(file_name, unit_details)


    '''print 'Updating PC for', unit_name, len(unit_details)
    base_unit_director = BASE_DIR + unit_name + SLASH

    # create directory for unit, if not exists
    if not os.path.exists(base_unit_director):
        os.makedirs(base_unit_director)

    pc_folder = base_unit_director + 'PC' + SLASH

    # create PC directory if not exists
    if not os.path.exists(pc_folder):
        os.makedirs(pc_folder)

    for j in range(len(table_list), len(unit_details)):
        table_list.append({})
        table_list[-1]['sr_no'] = ''
        table_list[-1]['description_of_goods'] = ''
        table_list[-1]['qty'] = ''
        table_list[-1]['cif_value'] = ''
        table_list[-1]['total'] = ''
        table_list[-1]['address'] = 'To,\n'
        table_list[-1]['short_address'] = '__________'
        table_list[-1]['balance'] = ''

    files_list = []
    for i in range(len(unit_details)):
        # print ''
        # print unit_details[i]
        # print '>>>', table_list[i]
        # print '------------------------------------------------------------------'

        temp_doc_file_name = 'temp' + str(i) + '.docx'
        save_to_pc_doc_file(temp_doc_file_name, unit_details[i], table_list[i])
        files_list.append(temp_doc_file_name)

    combine_word_documents(files_list, doc_file_name, EMPTY_FILE)

    # Re-read and save
    document = docx.Document(doc_file_name)
    document.save(doc_file_name)

    for file_name in files_list:
        os.remove(file_name)'''

def update_rwc_file(doc_file_name, unit_name, unit_details, table_list):
    print 'Updating RWC for', unit_name, len(unit_details)
    base_unit_director = BASE_DIR + unit_name + SLASH

    # create directory for unit, if not exists
    if not os.path.exists(base_unit_director):
        os.makedirs(base_unit_director)

    rwc_folder = base_unit_director + 'RWC' + SLASH

    # create RWC directory if not exists
    if not os.path.exists(rwc_folder):
        os.makedirs(rwc_folder)

    for j in range(len(table_list), len(unit_details)):
        table_list.append({})
        table_list[-1]['sr_no'] = ''
        table_list[-1]['description_of_goods'] = ''
        table_list[-1]['qty'] = ''
        table_list[-1]['cif_value'] = ''
        table_list[-1]['total'] = ''
        table_list[-1]['balance'] = ''

        table_list[-1]['particulars'] = ''
        table_list[-1]['description_of_goods'] = ''
        table_list[-1]['gross_weight'] = ''
        table_list[-1]['date_of_rewarehousing'] = ''
        table_list[-1]['shortage'] = ''
        table_list[-1]['address'] = 'To,\n'

    files_list = []
    for i in range(len(unit_details)):
        # print ''
        # print unit_details[i]
        # print '>>>', table_list[i]
        # print '------------------------------------------------------------------'

        temp_doc_file_name = 'temp' + str(i) + '.docx'
        save_to_rwc_doc_file(temp_doc_file_name, unit_details[i], table_list[i])
        files_list.append(temp_doc_file_name)

    combine_word_documents(files_list, doc_file_name, EMPTY_FILE)

    # Re-read and save
    document = Document(doc_file_name)
    document.save(doc_file_name)

    for file_name in files_list:
        os.remove(file_name)

def update_ic_file(unit_details_dict):
    base_unit_director = BASE_DIR
    
    # create directory for unit, if not exists
    if not os.path.exists(base_unit_director):
        os.makedirs(base_unit_director)

    ic_folder = base_unit_director + 'IC' + SLASH

    # create IC directory if not exists
    if not os.path.exists(ic_folder):
        os.makedirs(ic_folder)
    
    for unit_details in unit_details_dict:

        print unit_details
        file_name=ic_folder
        print  "+++++++++++++"
        print unit_details['ClientName']
        print unit_details['SubBranch'].strip()

        if not os.path.exists(ic_folder+unit_details['ClientName']+SLASH):
            os.makedirs(ic_folder+unit_details['ClientName']+SLASH)
        
        file_name = ic_folder+unit_details['ClientName']+SLASH
        if unit_details['SubBranch'].strip() != '':
            file_name = ic_folder+unit_details['ClientName']+SLASH+unit_details['SubBranch']+SLASH
            if not os.path.exists(ic_folder+unit_details['ClientName']+SLASH+unit_details['SubBranch']+SLASH):
                os.makedirs(ic_folder+unit_details['ClientName']+SLASH+unit_details['SubBranch']+SLASH)               
        else:
            file_name = ic_folder+unit_details['ClientName']+SLASH
      
        file_name = file_name + 'IC Application '

        file_name += (unit_details['ICApplicationDate'] + '({})'.format(unit_details['SerialNo'])).strip() + '.docx'
        print file_name            
        print "END+============================="       
        save_to_ic_doc_file(file_name, unit_details)

    '''dates_set = set()
    files_list = []
    for i in range(len(unit_details)):
        # print ''
        # print unit_details[i]
        # print '>>>', table_list[i]
        # print '------------------------------------------------------------------'

        file_name = ic_folder + 'IC Application '
        if unit_details[i]['ic_date'].strip() not in dates_set:
            dates_set.add(unit_details[i]['ic_date'].strip())
            file_name += (unit_details[i]['ic_date'] + '.docx').strip()
        else:
            no = 2
            while (unit_details[i]['ic_date'] + ' ({})'.format(no)).strip() in dates_set:
                no += 1
            dates_set.add((unit_details[i]['ic_date'] + ' ({})'.format(no)).strip())
            file_name += ((unit_details[i]['ic_date'] + ' ({})'.format(no)) + '.docx').strip()

        # print file_name
        table_dict = read_ic_doc_file(file_name)
        # print table_dict
        save_to_ic_doc_file(file_name, unit_details[i], table_dict)'''

def update(file_path, directory):
    global BASE_DIR, SLASH, EXCEL_FILE, TEMPLATES_DIR, PC_DOCX_TEMPLATE, IC_DOCX_TEMPLATE, EMPTY_FILE

    BASE_DIR = directory
    print BASE_DIR
    SLASH = '/'
    EXCEL_FILE =ntpath.basename(file_path)
    #print EXCEL_FILE
    print "==================================================================="
    TEMPLATES_DIR = '_files'
    PC_DOCX_TEMPLATE = TEMPLATES_DIR + SLASH + 'PC-Template - Marvel.docx'
    IC_DOCX_TEMPLATE = TEMPLATES_DIR + SLASH + 'IC-Template.docx'
    EMPTY_FILE = TEMPLATES_DIR + SLASH + 'empty.docx'

    unit_details_dict = view_excel_file(BASE_DIR, EXCEL_FILE)
    update_ic_file(unit_details_dict)
    update_pc_file(unit_details_dict)

    print "$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$$"
    '''for unit_name in unit_details:
        pc_doc_file_name = BASE_DIR + unit_name + SLASH + 'PC' + SLASH + 'PC ' + get_finiancial_year(datetime.datetime.now()) + '.docx'
        rwc_doc_file_name = BASE_DIR + unit_name + SLASH + 'RWC' + SLASH + 'RWC ' + get_finiancial_year(datetime.datetime.now()) + '.docx'

        pc_doc_file_contents = read_pc_doc_file(pc_doc_file_name)
        rwc_doc_file_contents = read_rwc_doc_file(rwc_doc_file_name)

        update_pc_file(pc_doc_file_name, unit_name, unit_details[unit_name], pc_doc_file_contents)
        update_ic_file(unit_name, unit_details[unit_name])
        update_rwc_file(rwc_doc_file_name, unit_name, unit_details[unit_name], rwc_doc_file_contents)'''

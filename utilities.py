import docx2txt
import docx
from docxtpl import *
from xml.dom.minidom import Text
from datetime import datetime

def get_finiancial_year(date):
    year = date.year
    if date.month < 4:
        year -= 1
    return str(year) + '-' + str(year+1)[2:]

def normalize_qty(qty):
    is_none = True
    a = 0
    for i in qty:
        number = (ord(i) - ord('0'))
        if number > 9 or number < 0:
            break
        a *= 10
        a += number
        is_none = False
    if not is_none:
        return str(a)
    else:
        return ''

def get_text_from_cell(cell):
    paragraphs = cell.paragraphs
    cell_text = ''
    for paragraph in paragraphs:
        cell_text += (paragraph.text.strip() + '\n')
    return cell_text

def combine_word_documents(files, combined_file_name, empty_file):
    EMPTY_FILE = empty_file
    empty_docx = docx.Document(EMPTY_FILE)
    empty_docx.save(EMPTY_FILE)
    combined_document = docx.Document(EMPTY_FILE)
    count, number_of_files = 0, len(files)

    for i in range(len(files)):
        combined_document.add_paragraph('{{p subdoc' + str(count) + '}}')
        if count < number_of_files - 1:
            combined_document.add_page_break()
        count += 1

    combined_document.save(combined_file_name)

    document_template = DocxTemplate(combined_file_name)
    count = 0
    context = {}
    for file in files:
        context['subdoc' + str(count)] = document_template.new_subdoc(file)
        count += 1
    document_template.render(context)
    document_template.save(combined_file_name)

def combine_word_documents_manual(files, combined_file_name, empty_file):
    EMPTY_FILE = empty_file
    empty_docx = docx.Document(EMPTY_FILE)
    empty_docx.save(EMPTY_FILE)
    combined_document = docx.Document(EMPTY_FILE)
    count, number_of_files = 0, len(files)
    for file in files:
        sub_doc = docx.Document(file)

        # Don't add a page break if you've
        # reached the last file.
        if count < number_of_files - 1:
            sub_doc.add_page_break()

        for element in sub_doc._part._element:
            combined_document._part._element.append(element)
        count += 1

    combined_document.save(combined_file_name)


def convert_to_xml(str_raw_text):
    tempText = Text()
    tempText.data = str_raw_text
    return tempText.toxml()

def check_expired():
    expiry_date = datetime.strptime('30/12/2016','%d/%m/%Y')
    now = datetime.now()
    return now > expiry_date

def remove_temp(xcel_list):
    clean_xcel_list = []
    for i in xcel_list:
        if i.find('~$') < 0 :
            clean_xcel_list.append(i.replace('\\','/'))
    return clean_xcel_list
	
def new_linify(string):
    string = string.replace('1.', '')
    i = 2
    while string.find(str(i) + ".") >= 0:
        string = string.replace(str(i) + ".", "\n")
        i += 1
    return string